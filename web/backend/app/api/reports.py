"""
Reports API endpoints for AI Reports functionality.
"""

import json
import logging
import io
from datetime import datetime
from typing import Optional, List
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.db.database import get_db
from app.db.models import User, Report, UserSettings
from app.auth.deps import get_current_active_user

logger = logging.getLogger(__name__)
router = APIRouter()


# =============================================================================
# Report Templates
# =============================================================================

REPORT_TEMPLATES = [
    {
        "id": "standard",
        "name": "Standard Report",
        "name_cn": "标准报告",
        "description": "General purpose report template",
        "icon": "file-text",
        "sections": [
            {"title": "执行摘要", "placeholder": "概述报告的主要发现和结论..."},
            {"title": "背景介绍", "placeholder": "描述报告的背景和目的..."},
            {"title": "主要发现", "placeholder": "列出主要的发现和分析结果..."},
            {"title": "详细分析", "placeholder": "提供详细的分析和支持数据..."},
            {"title": "结论与建议", "placeholder": "总结结论并提出建议..."},
        ],
    },
    {
        "id": "executive",
        "name": "Executive Summary",
        "name_cn": "高管摘要",
        "description": "Brief summary for executives",
        "icon": "briefcase",
        "sections": [
            {"title": "核心要点", "placeholder": "3-5个核心要点..."},
            {"title": "关键数据", "placeholder": "重要的数据指标..."},
            {"title": "行动建议", "placeholder": "需要采取的行动..."},
        ],
    },
    {
        "id": "technical",
        "name": "Technical Report",
        "name_cn": "技术报告",
        "description": "Detailed technical analysis",
        "icon": "code",
        "sections": [
            {"title": "技术概述", "placeholder": "技术方案的整体概述..."},
            {"title": "架构设计", "placeholder": "系统架构和设计说明..."},
            {"title": "实现细节", "placeholder": "关键实现细节和代码示例..."},
            {"title": "性能分析", "placeholder": "性能指标和优化建议..."},
            {"title": "测试结果", "placeholder": "测试方案和结果..."},
            {"title": "部署说明", "placeholder": "部署流程和注意事项..."},
        ],
    },
    {
        "id": "research",
        "name": "Research Report",
        "name_cn": "研究报告",
        "description": "In-depth research analysis",
        "icon": "search",
        "sections": [
            {"title": "研究背景", "placeholder": "研究的背景和动机..."},
            {"title": "文献综述", "placeholder": "相关文献和现有研究..."},
            {"title": "研究方法", "placeholder": "采用的研究方法..."},
            {"title": "数据分析", "placeholder": "数据收集和分析过程..."},
            {"title": "研究发现", "placeholder": "主要研究发现..."},
            {"title": "讨论", "placeholder": "结果讨论和意义解读..."},
            {"title": "结论", "placeholder": "研究结论..."},
            {"title": "参考文献", "placeholder": "引用的参考文献..."},
        ],
    },
]

# Chart type definitions
CHART_TYPES = [
    {"id": "line", "name": "折线图", "description": "展示趋势变化"},
    {"id": "bar", "name": "柱状图", "description": "比较不同类别"},
    {"id": "pie", "name": "饼图", "description": "展示占比分布"},
    {"id": "area", "name": "面积图", "description": "展示累积变化"},
]


# =============================================================================
# Request/Response Models
# =============================================================================

class SectionData(BaseModel):
    id: str
    title: str
    content: str = ""
    order: int = 0
    ai_generated: bool = False


class ChartDataPoint(BaseModel):
    name: str
    value: float
    extra: Optional[dict] = None


class ChartConfig(BaseModel):
    xKey: str = "name"
    yKey: str = "value"
    colors: List[str] = ["#3B82F6", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6"]


class ChartData(BaseModel):
    id: str
    type: str  # line, bar, pie, area
    title: str
    data: List[ChartDataPoint]
    config: Optional[ChartConfig] = None


class ReportCreate(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    template_type: str = "standard"
    research_session_id: Optional[str] = None


class ReportUpdate(BaseModel):
    title: Optional[str] = Field(None, max_length=200)
    content: Optional[str] = None
    content_html: Optional[str] = None
    status: Optional[str] = None


class SectionUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None


class ChartCreate(BaseModel):
    type: str = Field(..., pattern="^(line|bar|pie|area)$")
    title: str = Field(..., min_length=1, max_length=100)
    data: List[ChartDataPoint]
    config: Optional[ChartConfig] = None


class ChartUpdate(BaseModel):
    title: Optional[str] = None
    data: Optional[List[ChartDataPoint]] = None
    config: Optional[ChartConfig] = None


class ReportResponse(BaseModel):
    id: str
    title: str
    template_type: str
    status: str
    sections: List[dict]
    charts_data: List[dict]
    content: str
    content_html: str
    export_format: Optional[str]
    exported_at: Optional[datetime]
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class ReportListResponse(BaseModel):
    reports: List[ReportResponse]
    total: int
    page: int
    page_size: int


class TemplateResponse(BaseModel):
    id: str
    name: str
    name_cn: str
    description: str
    icon: str
    sections: List[dict]


class TemplateListResponse(BaseModel):
    templates: List[TemplateResponse]


class AIGenerateRequest(BaseModel):
    section_id: Optional[str] = None  # If None, generate entire report
    prompt: Optional[str] = None  # Additional instructions
    model: str = "openai/gpt-4o"


class ExportRequest(BaseModel):
    format: str = Field(..., pattern="^(pdf|docx|md)$")


# =============================================================================
# Helper Functions
# =============================================================================

def build_report_prompt(report: Report, section_id: Optional[str] = None, prompt: Optional[str] = None) -> tuple:
    """Build system and user prompts for AI report generation."""

    sections = report.sections or []
    section_titles = [s.get("title", "") for s in sections]

    if section_id:
        # Generate specific section
        section = next((s for s in sections if s.get("id") == section_id), None)
        if not section:
            return None, None

        section_title = section.get("title", "")
        existing_content = section.get("content", "")

        system_prompt = f"""你是一位专业的报告撰写助手。请为报告《{report.title}》撰写"{section_title}"章节的内容。

报告类型：{report.template_type}
报告结构：{', '.join(section_titles)}

要求：
1. 内容专业、结构清晰
2. 使用 Markdown 格式
3. 如有数据或分析，请提供具体的论述
4. 直接输出章节内容，不要添加章节标题（标题已存在）"""

        user_prompt = f'请撰写「{section_title}」章节。'
        if existing_content:
            user_prompt += f"\n\n现有内容（可参考或扩展）：\n{existing_content[:500]}"
        if prompt:
            user_prompt += f"\n\n额外要求：{prompt}"

    else:
        # Generate entire report
        system_prompt = f"""你是一位专业的报告撰写助手。请为报告《{report.title}》生成完整的内容。

报告类型：{report.template_type}
报告结构：
{chr(10).join([f'- {t}' for t in section_titles])}

要求：
1. 按照给定的章节结构逐一撰写
2. 每个章节使用二级标题（## 章节名）
3. 内容专业、有深度
4. 使用 Markdown 格式
5. 如需展示数据，可以使用表格"""

        user_prompt = "请生成完整的报告内容。"
        if prompt:
            user_prompt += f"\n\n额外要求：{prompt}"

    return system_prompt, user_prompt


def generate_pdf(report: Report) -> bytes:
    """Generate PDF from report content using WeasyPrint."""
    try:
        from weasyprint import HTML, CSS
        import markdown

        # Convert markdown to HTML
        md_content = report.content or ""
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

        # Create full HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{report.title}</title>
            <style>
                body {{
                    font-family: 'Helvetica Neue', Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 40px;
                }}
                h1 {{ color: #1a1a1a; border-bottom: 2px solid #3B82F6; padding-bottom: 10px; }}
                h2 {{ color: #2d2d2d; margin-top: 30px; }}
                h3 {{ color: #4a4a4a; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
                th {{ background-color: #f5f5f5; }}
                code {{ background-color: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                blockquote {{ border-left: 4px solid #3B82F6; margin: 20px 0; padding-left: 20px; color: #666; }}
            </style>
        </head>
        <body>
            <h1>{report.title}</h1>
            {html_content}
        </body>
        </html>
        """

        # Generate PDF
        html = HTML(string=full_html)
        pdf_bytes = html.write_pdf()

        return pdf_bytes
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


def generate_docx(report: Report) -> bytes:
    """Generate DOCX from report content using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # Title
        title = doc.add_heading(report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Content - simple markdown to docx conversion
        content = report.content or ""
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Handle headings
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            else:
                doc.add_paragraph(line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


def generate_markdown(report: Report) -> bytes:
    """Generate Markdown export."""
    content = f"# {report.title}\n\n"
    content += report.content or ""
    return content.encode('utf-8')


# =============================================================================
# Endpoints
# =============================================================================

@router.get("/templates", response_model=TemplateListResponse)
async def get_templates():
    """Get all available report templates."""
    return TemplateListResponse(
        templates=[
            TemplateResponse(**template) for template in REPORT_TEMPLATES
        ]
    )


@router.get("/chart-types")
async def get_chart_types():
    """Get available chart types."""
    return {"chart_types": CHART_TYPES}


@router.get("/reports", response_model=ReportListResponse)
async def list_reports(
    page: int = 1,
    page_size: int = 20,
    status: Optional[str] = None,
    template_type: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """List all reports for the current user."""
    query = (
        db.query(Report)
        .filter(Report.user_id == current_user.id)
    )

    if status:
        query = query.filter(Report.status == status)
    if template_type:
        query = query.filter(Report.template_type == template_type)

    query = query.order_by(Report.updated_at.desc())

    total = query.count()
    reports = query.offset((page - 1) * page_size).limit(page_size).all()

    return ReportListResponse(
        reports=[
            ReportResponse(
                id=r.id,
                title=r.title,
                template_type=r.template_type,
                status=r.status or "draft",
                sections=r.sections or [],
                charts_data=r.charts_data or [],
                content=r.content or "",
                content_html=r.content_html or "",
                export_format=r.export_format,
                exported_at=r.exported_at,
                created_at=r.created_at,
                updated_at=r.updated_at,
            )
            for r in reports
        ],
        total=total,
        page=page,
        page_size=page_size,
    )


@router.post("/reports", response_model=ReportResponse, status_code=status.HTTP_201_CREATED)
async def create_report(
    request: ReportCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Create a new report."""
    # Get template sections
    template = next(
        (t for t in REPORT_TEMPLATES if t["id"] == request.template_type),
        REPORT_TEMPLATES[0]
    )

    # Initialize sections from template
    sections = [
        {
            "id": str(uuid4()),
            "title": s["title"],
            "content": "",
            "order": i,
            "ai_generated": False,
        }
        for i, s in enumerate(template.get("sections", []))
    ]

    report = Report(
        id=str(uuid4()),
        user_id=current_user.id,
        research_session_id=request.research_session_id,
        title=request.title,
        template_type=request.template_type,
        sections=sections,
        charts_data=[],
        status="draft",
    )
    db.add(report)
    db.commit()
    db.refresh(report)

    return ReportResponse(
        id=report.id,
        title=report.title,
        template_type=report.template_type,
        status=report.status or "draft",
        sections=report.sections or [],
        charts_data=report.charts_data or [],
        content=report.content or "",
        content_html=report.content_html or "",
        export_format=report.export_format,
        exported_at=report.exported_at,
        created_at=report.created_at,
        updated_at=report.updated_at,
    )


@router.get("/reports/{report_id}", response_model=ReportResponse)
async def get_report(
    report_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Get a report by ID."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    return ReportResponse(
        id=report.id,
        title=report.title,
        template_type=report.template_type,
        status=report.status or "draft",
        sections=report.sections or [],
        charts_data=report.charts_data or [],
        content=report.content or "",
        content_html=report.content_html or "",
        export_format=report.export_format,
        exported_at=report.exported_at,
        created_at=report.created_at,
        updated_at=report.updated_at,
    )


@router.put("/reports/{report_id}", response_model=ReportResponse)
async def update_report(
    report_id: str,
    request: ReportUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Update a report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    if request.title is not None:
        report.title = request.title
    if request.content is not None:
        report.content = request.content
    if request.content_html is not None:
        report.content_html = request.content_html
    if request.status is not None:
        report.status = request.status

    report.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(report)

    return ReportResponse(
        id=report.id,
        title=report.title,
        template_type=report.template_type,
        status=report.status or "draft",
        sections=report.sections or [],
        charts_data=report.charts_data or [],
        content=report.content or "",
        content_html=report.content_html or "",
        export_format=report.export_format,
        exported_at=report.exported_at,
        created_at=report.created_at,
        updated_at=report.updated_at,
    )


@router.delete("/reports/{report_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_report(
    report_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Delete a report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    db.delete(report)
    db.commit()


# =============================================================================
# Section Endpoints
# =============================================================================

@router.put("/reports/{report_id}/sections/{section_id}")
async def update_section(
    report_id: str,
    section_id: str,
    request: SectionUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Update a specific section."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    sections = report.sections or []
    section_found = False

    for section in sections:
        if section.get("id") == section_id:
            if request.title is not None:
                section["title"] = request.title
            if request.content is not None:
                section["content"] = request.content
            section_found = True
            break

    if not section_found:
        raise HTTPException(status_code=404, detail="Section not found")

    report.sections = sections
    report.updated_at = datetime.utcnow()

    # Update main content from sections
    report.content = "\n\n".join([
        f"## {s.get('title', '')}\n\n{s.get('content', '')}"
        for s in sections if s.get('content')
    ])

    db.commit()

    return {"message": "Section updated", "section_id": section_id}


@router.post("/reports/{report_id}/sections")
async def add_section(
    report_id: str,
    request: SectionUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Add a new section to the report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    sections = report.sections or []
    new_section = {
        "id": str(uuid4()),
        "title": request.title or "新章节",
        "content": request.content or "",
        "order": len(sections),
        "ai_generated": False,
    }
    sections.append(new_section)

    report.sections = sections
    report.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Section added", "section": new_section}


@router.delete("/reports/{report_id}/sections/{section_id}")
async def delete_section(
    report_id: str,
    section_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Delete a section from the report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    sections = report.sections or []
    sections = [s for s in sections if s.get("id") != section_id]

    # Reorder remaining sections
    for i, section in enumerate(sections):
        section["order"] = i

    report.sections = sections
    report.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Section deleted"}


# =============================================================================
# Chart Endpoints
# =============================================================================

@router.post("/reports/{report_id}/charts")
async def add_chart(
    report_id: str,
    request: ChartCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Add a chart to the report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    charts = report.charts_data or []
    new_chart = {
        "id": str(uuid4()),
        "type": request.type,
        "title": request.title,
        "data": [d.dict() for d in request.data],
        "config": request.config.dict() if request.config else {
            "xKey": "name",
            "yKey": "value",
            "colors": ["#3B82F6", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6"]
        },
    }
    charts.append(new_chart)

    report.charts_data = charts
    report.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Chart added", "chart": new_chart}


@router.put("/reports/{report_id}/charts/{chart_id}")
async def update_chart(
    report_id: str,
    chart_id: str,
    request: ChartUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Update a chart."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    charts = report.charts_data or []
    chart_found = False

    for chart in charts:
        if chart.get("id") == chart_id:
            if request.title is not None:
                chart["title"] = request.title
            if request.data is not None:
                chart["data"] = [d.dict() for d in request.data]
            if request.config is not None:
                chart["config"] = request.config.dict()
            chart_found = True
            break

    if not chart_found:
        raise HTTPException(status_code=404, detail="Chart not found")

    report.charts_data = charts
    report.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Chart updated", "chart_id": chart_id}


@router.delete("/reports/{report_id}/charts/{chart_id}")
async def delete_chart(
    report_id: str,
    chart_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Delete a chart from the report."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    charts = report.charts_data or []
    charts = [c for c in charts if c.get("id") != chart_id]

    report.charts_data = charts
    report.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Chart deleted"}


# =============================================================================
# AI Generation Endpoint
# =============================================================================

@router.post("/reports/{report_id}/generate")
async def generate_report_content(
    report_id: str,
    request: AIGenerateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Generate report content using AI with streaming response."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    async def generate_response():
        try:
            # Get user's API settings
            settings = db.query(UserSettings).filter(
                UserSettings.user_id == current_user.id
            ).first()

            # Build prompts
            system_prompt, user_prompt = build_report_prompt(
                report=report,
                section_id=request.section_id,
                prompt=request.prompt,
            )

            if not system_prompt:
                yield f"data: {json.dumps({'error': 'Section not found'})}\n\n"
                return

            # Get API key based on model
            model = request.model or "openai/gpt-4o"
            api_key = None

            if model.startswith("openai/"):
                api_key = settings.openai_api_key if settings else None
            elif model.startswith("anthropic/"):
                api_key = settings.anthropic_api_key if settings else None
            elif model.startswith("google/"):
                api_key = settings.google_api_key if settings else None

            if not api_key:
                error_msg = "请先在设置中配置相应的 API 密钥"
                yield f"data: {json.dumps({'error': error_msg})}\n\n"
                return

            full_response = ""

            if model.startswith("openai/"):
                import openai
                client = openai.OpenAI(api_key=api_key)
                model_name = model.replace("openai/", "")

                stream = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    stream=True,
                )

                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_response += content
                        yield f"data: {json.dumps({'content': content})}\n\n"

            elif model.startswith("anthropic/"):
                import anthropic
                client = anthropic.Anthropic(api_key=api_key)
                model_name = model.replace("anthropic/", "")

                with client.messages.stream(
                    model=model_name,
                    max_tokens=8192,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}],
                ) as stream:
                    for text in stream.text_stream:
                        full_response += text
                        yield f"data: {json.dumps({'content': text})}\n\n"

            elif model.startswith("google/"):
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model_name = model.replace("google/", "")

                gemini = genai.GenerativeModel(
                    model_name,
                    system_instruction=system_prompt,
                )
                response = gemini.generate_content(user_prompt, stream=True)

                for chunk in response:
                    if chunk.text:
                        full_response += chunk.text
                        yield f"data: {json.dumps({'content': chunk.text})}\n\n"

            # Update report/section with generated content
            if request.section_id:
                sections = report.sections or []
                for section in sections:
                    if section.get("id") == request.section_id:
                        section["content"] = full_response
                        section["ai_generated"] = True
                        break
                report.sections = sections
            else:
                report.content = full_response

            report.updated_at = datetime.utcnow()
            db.commit()

            yield f"data: {json.dumps({'done': True, 'section_id': request.section_id})}\n\n"

        except Exception as e:
            logger.error(f"AI generation error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(
        generate_response(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        },
    )


# =============================================================================
# Export Endpoint
# =============================================================================

@router.post("/reports/{report_id}/export")
async def export_report(
    report_id: str,
    request: ExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export report to PDF, DOCX, or Markdown."""
    report = (
        db.query(Report)
        .filter(
            Report.id == report_id,
            Report.user_id == current_user.id,
        )
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    # Generate content from sections if main content is empty
    if not report.content and report.sections:
        report.content = "\n\n".join([
            f"## {s.get('title', '')}\n\n{s.get('content', '')}"
            for s in report.sections if s.get('content')
        ])

    # Generate export based on format
    if request.format == "pdf":
        file_bytes = generate_pdf(report)
        media_type = "application/pdf"
        filename = f"{report.title}.pdf"
    elif request.format == "docx":
        file_bytes = generate_docx(report)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{report.title}.docx"
    else:  # markdown
        file_bytes = generate_markdown(report)
        media_type = "text/markdown"
        filename = f"{report.title}.md"

    # Update export info
    report.export_format = request.format
    report.exported_at = datetime.utcnow()
    db.commit()

    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )
