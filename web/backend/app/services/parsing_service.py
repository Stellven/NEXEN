"""
Document parsing service for extracting text from various file formats.
"""

import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)


class ParsingService:
    """Service for parsing documents and extracting text content."""

    async def parse_file(self, file_path: str, file_type: str) -> str:
        """Parse a file and extract text content."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower()

        if file_type == "pdf":
            return await self.parse_pdf(file_path)
        elif file_type == "docx":
            return await self.parse_docx(file_path)
        elif file_type in ("md", "markdown"):
            return await self.parse_text(file_path)
        elif file_type == "txt":
            return await self.parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def parse_pdf(self, file_path: str) -> str:
        """Parse PDF file using PyMuPDF."""
        try:
            import fitz  # PyMuPDF

            text_parts = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)

            content = "\n\n".join(text_parts)
            logger.info(f"Parsed PDF: {file_path}, {len(content)} chars")
            return content

        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            raise ImportError("PyMuPDF is required for PDF parsing")
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            raise

    async def parse_docx(self, file_path: str) -> str:
        """Parse DOCX file using python-docx."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            content = "\n\n".join(text_parts)
            logger.info(f"Parsed DOCX: {file_path}, {len(content)} chars")
            return content

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX parsing")
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            raise

    async def parse_text(self, file_path: str) -> str:
        """Parse plain text or markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            logger.info(f"Parsed text file: {file_path}, {len(content)} chars")
            return content

        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"Text parse error: {e}")
            raise

    async def parse_url(self, url: str) -> str:
        """Extract content from URL using trafilatura."""
        try:
            import trafilatura

            # Download and extract
            downloaded = trafilatura.fetch_url(url)
            if not downloaded:
                raise ValueError(f"Failed to fetch URL: {url}")

            content = trafilatura.extract(
                downloaded,
                include_comments=False,
                include_tables=True,
                include_images=False,
                include_links=False,
                output_format="txt",
            )

            if not content:
                raise ValueError(f"Failed to extract content from URL: {url}")

            logger.info(f"Parsed URL: {url}, {len(content)} chars")
            return content

        except ImportError:
            logger.error("trafilatura not installed. Install with: pip install trafilatura")
            raise ImportError("trafilatura is required for URL parsing")
        except Exception as e:
            logger.error(f"URL parse error: {e}")
            raise


class TextChunker:
    """Utility for splitting text into chunks with overlap."""

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk_by_words(self, text: str) -> list[str]:
        """Split text into chunks by word count."""
        if not text:
            return []

        words = text.split()
        chunks = []
        start = 0

        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start += self.chunk_size - self.overlap

        return chunks

    def chunk_by_tokens(self, text: str) -> list[str]:
        """Split text into chunks by token count using tiktoken."""
        try:
            import tiktoken

            encoding = tiktoken.get_encoding("cl100k_base")
            tokens = encoding.encode(text)
            chunks = []
            start = 0

            while start < len(tokens):
                end = min(start + self.chunk_size, len(tokens))
                chunk_tokens = tokens[start:end]
                chunk = encoding.decode(chunk_tokens)
                chunks.append(chunk)
                start += self.chunk_size - self.overlap

            return chunks

        except ImportError:
            logger.warning("tiktoken not installed, falling back to word-based chunking")
            return self.chunk_by_words(text)

    def chunk_by_paragraphs(self, text: str, max_chunk_size: int = 2000) -> list[str]:
        """Split text into chunks by paragraphs, respecting max size."""
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_size = len(para.split())

            if current_size + para_size > max_chunk_size and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_size = 0

            current_chunk.append(para)
            current_size += para_size

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks
